import os
import io
import json
import math
from datetime import datetime
from typing import Dict, Any, List, Tuple

import numpy as np
import numpy_financial as npf
import pandas as pd
import plotly.express as px
import requests
import streamlit as st

# --- Exports ---
from docx import Document
from docx.shared import Pt
import markdown2
from xhtml2pdf import pisa


# =========================
# ----- HELPER LOGIC -----
# =========================

def annuity_payment(principal: float, annual_rate: float, years: int) -> float:
    """Equal annual debt service (principal + interest)."""
    if annual_rate == 0:
        return principal / years
    r = annual_rate
    n = years
    return principal * (r * (1 + r) ** n) / ((1 + r) ** n - 1)


def build_debt_schedule(
    principal: float, annual_rate: float, years: int
) -> pd.DataFrame:
    """Amortization schedule with equal total payment."""
    pay = annuity_payment(principal, annual_rate, years)
    rows = []
    bal = principal
    for yr in range(1, years + 1):
        interest = bal * annual_rate
        principal_pay = max(0.0, pay - interest)
        bal = max(0.0, bal - principal_pay)
        rows.append(
            {
                "year": yr,
                "debt_opening": round(bal + principal_pay, 2),
                "debt_payment_total": round(pay, 2),
                "interest": round(interest, 2),
                "principal": round(principal_pay, 2),
                "debt_closing": round(bal, 2),
            }
        )
    return pd.DataFrame(rows)


def compute_financials(assump: Dict[str, Any]) -> Dict[str, Any]:
    """Core 10y model + simple DSCR/IRR."""
    years = int(assump.get("years", 10))
    y1_rev = float(assump.get("year1_revenue", 450_000_000))
    growth = float(assump.get("revenue_growth", 0.08))
    ebitda_margin = float(assump.get("ebitda_margin", 0.22))
    capex_total = float(assump.get("capex_total", 1_200_000_000))
    interest_rate = float(assump.get("interest_rate", 0.11))
    tenor_years = int(assump.get("tenor_years", 7))
    debt_ratio = float(assump.get("debt_ratio", 0.7))
    tax_rate = float(assump.get("tax_rate", 0.25))  # simple flat tax on EBIT

    # Revenue & EBITDA
    revenue = [y1_rev * ((1 + growth) ** i) for i in range(years)]
    ebitda = [r * ebitda_margin for r in revenue]

    # Assume D&A ~ 5% of CAPEX straight-line across 10 years (very rough)
    da = [capex_total * 0.05] * years

    # Debt schedule
    debt_amt = capex_total * debt_ratio
    debt = build_debt_schedule(debt_amt, interest_rate, tenor_years)
    # Extend to 10 years with zeroes once loan is repaid
    if tenor_years < years:
        tail = pd.DataFrame(
            {
                "year": list(range(tenor_years + 1, years + 1)),
                "debt_opening": 0.0,
                "debt_payment_total": 0.0,
                "interest": 0.0,
                "principal": 0.0,
                "debt_closing": 0.0,
            }
        )
        debt = pd.concat([debt, tail], ignore_index=True)

    # EBIT, tax, FCF (very simplified)
    ebit = [max(0.0, ebitda[i] - da[i]) for i in range(years)]
    taxes = [ebit[i] * tax_rate for i in range(years)]
    # Operating CF before debt service (no WC here)
    ocf = [ebitda[i] - taxes[i] for i in range(years)]

    # DSCR = OCF / Debt Service (yrs with zero debt service -> NaN)
    ds = debt["debt_payment_total"].to_list()
    dscr = [ (ocf[i] / ds[i] if ds[i] > 0 else np.nan) for i in range(years) ]

    # IRR on Project (equity view simple):
    equity = capex_total * (1 - debt_ratio)
    # Equity cash flows: -equity upfront (t0), then OCF - debt service for years
    eq_cf = [-equity] + [ocf[i] - ds[i] for i in range(years)]
    try:
        irr_equity = float(npf.irr(eq_cf))
    except Exception:
        irr_equity = float("nan")

    # Payback (on equity): cumulative >= 0
    cum = 0.0
    payback_year = None
    for idx, cf in enumerate(eq_cf[1:], start=1):
        cum += cf
        if cum >= 0 and payback_year is None:
            payback_year = idx

    # Build DataFrame summary per year
    df = pd.DataFrame(
        {
            "Year": list(range(1, years + 1)),
            "Revenue": revenue,
            "EBITDA": ebitda,
            "D&A": da,
            "EBIT": ebit,
            "Tax": taxes,
            "OCF_before_DebtSvc": ocf,
            "DebtService": ds[:years],
            "DSCR": dscr[:years],
            "DebtClosing": debt["debt_closing"][:years].to_list(),
        }
    )

    metrics = {
        "revenue_yearly": revenue,
        "ebitda_yearly": ebitda,
        "ebitda_margin": ebitda_margin,
        "irr_equity": irr_equity,
        "dscr_min": float(np.nanmin(dscr)),
        "dscr_avg": float(np.nanmean([x for x in dscr if not math.isnan(x)]))
            if any([not math.isnan(x) for x in dscr]) else float("nan"),
        "payback_year": payback_year,
        "debt_ratio": debt_ratio,
        "tax_rate": tax_rate,
    }
    return {"df": df, "debt": debt, "metrics": metrics}


def default_prompt() -> str:
    return (
        "Generate a detailed investor-grade business plan for a 200 TPD plant protein "
        "facility in India. Use the provided metrics and assumptions. Write 60–100 pages worth of "
        "sections when expanded to Word/PDF, but return concise, well-structured Markdown with rich headings, "
        "tables and annexures I can later export. Include Executive Summary, Market & Competitive Landscape, "
        "Technical Feasibility, CAPEX/OPEX, Financials, Risks, Execution Roadmap, and Appendices."
    )


def build_markdown_report(assump: Dict[str, Any], metrics: Dict[str, Any]) -> str:
    # A crisp, deterministic scaffold the user can expand (GPT can overwrite this)
    lines = []
    lines.append("# Business Plan Report\n")
    lines.append("## 1. Executive Summary")
    lines.append(f"- Total CAPEX (input): ₹{assump['capex_total']:,}")
    lines.append(f"- EBITDA margin: {metrics['ebitda_margin']:.2%}")
    lines.append(f"- Equity IRR (simple): {metrics['irr_equity']:.2%}")
    lines.append(
        f"- DSCR: min {metrics['dscr_min']:.2f} | avg {metrics['dscr_avg']:.2f}"
    )
    lines.append(f"- Payback (years, simple): {metrics['payback_year']}\n")

    lines.append("## 2. Assumptions")
    lines.append("```json")
    lines.append(json.dumps(assump, indent=2))
    lines.append("```\n")

    lines.append("## 3. Financial KPI Highlights")
    lines.append("| KPI | Value |")
    lines.append("|---|---|")
    lines.append(f"| EBITDA margin | {metrics['ebitda_margin']:.2%} |")
    lines.append(f"| Equity IRR | {metrics['irr_equity']:.2%} |")
    lines.append(f"| DSCR (min) | {metrics['dscr_min']:.2f} |")
    lines.append(f"| DSCR (avg) | {metrics['dscr_avg']:.2f} |")
    lines.append(f"| Payback (years) | {metrics['payback_year']} |\n")

    lines.append("## 4. Sections (placeholders)")
    for sec in [
        "Promoter Profile", "Industry & Market Analysis", "Technical Feasibility",
        "Location & Infrastructure", "Project Costs (CAPEX)", "Operating Costs (OPEX)",
        "Revenue & Financial Analysis", "Compliance & Incentives",
        "Risk Register & Mitigation", "Execution Plan", "Annexures"
    ]:
        lines.append(f"### {sec}\n_TBD—use **Generate with GPT** to auto-write this section._\n")

    return "\n".join(lines)


def markdown_to_pdf_bytes(md_text: str) -> bytes:
    """Convert Markdown -> HTML -> PDF (bytes) using markdown2 + xhtml2pdf."""
    html = markdown2.markdown(md_text, extras=["tables", "fenced-code-blocks", "footnotes"])
    # Simple stylesheet for nicer PDF
    style = """
    <style>
    body { font-family: DejaVu Sans, Helvetica, Arial, sans-serif; font-size: 12px; }
    h1,h2,h3 { color: #222; }
    table { border-collapse: collapse; width: 100%; }
    th, td { border: 1px solid #999; padding: 6px; }
    code { background: #f2f2f2; padding: 2px 4px; }
    </style>
    """
    html = f"<!doctype html><html><head><meta charset='utf-8'>{style}</head><body>{html}</body></html>"
    out = io.BytesIO()
    pisa_status = pisa.CreatePDF(src=io.StringIO(html), dest=out)
    if pisa_status.err:
        raise RuntimeError("PDF generation failed.")
    return out.getvalue()


def markdown_to_docx_bytes(md_text: str) -> bytes:
    """Very simple Markdown -> DOCX; preserves headings and paragraphs."""
    # (For richer md->docx use 'mammoth' with HTML roundtrip. This keeps deps light.)
    doc = Document()
    styles = doc.styles['Normal']
    styles.font.name = 'Calibri'
    styles.font.size = Pt(11)

    for line in md_text.splitlines():
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
        else:
            doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def to_csv_bytes(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8")


def call_local_api(base_url: str, payload: Dict[str, Any], timeout: int = 120) -> Dict[str, Any]:
    url = base_url.rstrip("/") + "/generate"
    r = requests.post(url, json=payload, timeout=timeout)
    r.raise_for_status()
    return r.json()


def call_openai_gpt(api_key: str, model: str, sys_prompt: str, user_prompt: str) -> str:
    """
    Uses OpenAI's Python SDK to create a Chat Completion and return Markdown text.
    """
    # Lazy import so the app still runs without the package until user enables GPT.
    from openai import OpenAI
    client = OpenAI(api_key=api_key)

    # Chat Completions API (Python SDK)
    # Ref: OpenAI docs show `client.chat.completions.create(model="gpt-4o-mini", messages=[...])`.
    # See examples on platform docs/community. :contentReference[oaicite:0]{index=0}
    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.3,
    )
    return resp.choices[0].message.content


# =========================
# ----- STREAMLIT UI ------
# =========================

st.set_page_config(
    page_title="BizPlan Generator",
    page_icon="📈",
    layout="wide",
    menu_items={"about": "BizPlan Streamlit — generate, visualize, and export investor-grade plans."},
)

st.title("📈 BizPlan — Plant Protein Facility (200 TPD)")
st.caption("Generate a full investor-grade plan, visualize metrics, and export as PDF/DOCX/CSV/Markdown.")

# ----- Sidebar: Inputs -----
st.sidebar.header("Assumptions")
with st.sidebar.form("assumptions_form"):
    capex_total = st.number_input("Total CAPEX (₹)", 0, step=10_000_000, value=1_200_000_000)
    years = st.number_input("Model Years", 1, 50, 10)
    y1_rev = st.number_input("Year 1 Revenue (₹)", 0, step=10_000_000, value=450_000_000)
    growth = st.number_input("Revenue Growth (YoY, %)", 0.0, 1.0, 0.08, step=0.01, format="%.2f")
    ebitda_margin = st.number_input("EBITDA Margin (%)", 0.0, 1.0, 0.22, step=0.01, format="%.2f")
    interest_rate = st.number_input("Interest Rate (%)", 0.0, 1.0, 0.11, step=0.005, format="%.3f")
    tenor_years = st.number_input("Debt Tenor (yrs)", 1, 50, 7)
    debt_ratio = st.number_input("Debt Ratio", 0.0, 1.0, 0.70, step=0.05)
    tax_rate = st.number_input("Tax Rate", 0.0, 1.0, 0.25, step=0.01)
    submitted = st.form_submit_button("Recalculate")

assump = {
    "capex_total": capex_total,
    "years": years,
    "year1_revenue": y1_rev,
    "revenue_growth": growth,
    "ebitda_margin": ebitda_margin,
    "interest_rate": interest_rate,
    "tenor_years": tenor_years,
    "debt_ratio": debt_ratio,
    "tax_rate": tax_rate,
}

# --- Compute model ---
calc = compute_financials(assump)
df, debt_df, metrics = calc["df"], calc["debt"], calc["metrics"]

# --- Session state for markdown ---
if "report_md" not in st.session_state:
    st.session_state["report_md"] = build_markdown_report(assump, metrics)

# ----- Tabs -----
tab1, tab2, tab3, tab4 = st.tabs(
    ["📊 Dashboard", "🧠 Generate with GPT", "🧾 Report (Markdown)", "⚙️ Integrations & Export"]
)

# ========== TAB 1: Dashboard ==========
with tab1:
    col_a, col_b, col_c, col_d = st.columns(4)
    col_a.metric("Equity IRR (simple)", f"{metrics['irr_equity']*100:,.2f}%")
    col_b.metric("DSCR (min)", f"{metrics['dscr_min']:.2f}")
    col_c.metric("DSCR (avg)", f"{metrics['dscr_avg']:.2f}")
    col_d.metric("Payback (yrs)", metrics['payback_year'] if metrics['payback_year'] else "n/a")

    c1, c2 = st.columns(2)
    with c1:
        st.subheader("Revenue (₹) & EBITDA (₹)")
        df_long = df.melt(id_vars="Year", value_vars=["Revenue", "EBITDA"], var_name="Item", value_name="Amount")
        fig1 = px.line(df_long, x="Year", y="Amount", color="Item")
        st.plotly_chart(fig1, use_container_width=True)
    with c2:
        st.subheader("DSCR by Year")
        fig2 = px.bar(df, x="Year", y="DSCR")
        st.plotly_chart(fig2, use_container_width=True)

    st.subheader("Debt Amortization Schedule")
    st.dataframe(debt_df, use_container_width=True)

    st.subheader("Financial Table")
    st.dataframe(df, use_container_width=True)

    # Sensitivity (growth vs margin)
    st.markdown("### Sensitivity — EBITDA Margin vs Revenue Growth")
    g_vals = np.linspace(max(0, growth - 0.05), growth + 0.05, 5)
    m_vals = np.linspace(max(0, ebitda_margin - 0.05), min(0.6, ebitda_margin + 0.05), 5)
    grid = []
    for g in g_vals:
        for m in m_vals:
            tmp = compute_financials({**assump, "revenue_growth": float(g), "ebitda_margin": float(m)})
            grid.append({"Growth": g, "Margin": m, "Eq_IRR": tmp["metrics"]["irr_equity"]})
    heat = pd.DataFrame(grid)
    fig3 = px.density_heatmap(
        heat, x=heat["Growth"].round(3), y=heat["Margin"].round(3),
        z=heat["Eq_IRR"], histfunc="avg", text_auto=True, nbinsx=5, nbinsy=5
    ).update_layout(xaxis_title="Rev Growth", yaxis_title="EBITDA Margin", coloraxis_colorbar_title="Eq IRR")
    st.plotly_chart(fig3, use_container_width=True)


# ========== TAB 2: GPT ==========
with tab2:
    st.markdown("Use GPT to auto-draft or refine the **full report** based on your assumptions & metrics.")

    api_key = st.text_input("OpenAI API Key (stored only in this session)", value=os.getenv("OPENAI_API_KEY", ""), type="password")
    model = st.selectbox("Model", ["gpt-4o-mini", "gpt-4o", "gpt-4.1-mini", "gpt-4.1"])
    system_prompt = st.text_area(
        "System Prompt",
        "You are an expert investment banker and process engineer creating investor-grade business plans in concise Markdown.",
        height=80
    )
    user_prompt = st.text_area("User Prompt", default_prompt(), height=140)

    enhance_with_tables = st.checkbox("Ask GPT to expand detailed tables/annexures", value=True)
    tone = st.selectbox("Tone", ["Neutral professional", "Investor bullish", "Risk-balanced", "Technical deep-dive"])

    colg1, colg2 = st.columns([1,1])
    with colg1:
        if st.button("🚀 Generate Report with GPT", use_container_width=True, type="primary", disabled=not api_key):
            seed = {
                "assumptions": assump,
                "metrics": metrics,
                "now": datetime.utcnow().isoformat()
            }
            enriched_prompt = (
                f"{user_prompt}\n\nHere are the numeric inputs as JSON:\n```json\n{json.dumps(seed, indent=2)}\n```\n"
                f"Tone: {tone}. "
                f"{'Include rich tables (CAPEX/OPEX), risk register, compliance matrix, gantt-like roadmap.' if enhance_with_tables else ''}"
            )
            try:
                md = call_openai_gpt(api_key, model, system_prompt, enriched_prompt)
                st.session_state["report_md"] = md
                st.success("GPT report generated.")
            except Exception as e:
                st.error(f"OpenAI call failed: {e}")

    with colg2:
        st.write(" ")
        st.caption("Tip: You can also keep the scaffolded report in Tab 3 and only use GPT for specific sections.")

# ========== TAB 3: Report (Markdown) ==========
with tab3:
    st.markdown("Edit or paste your Markdown. Export as **PDF** or **DOCX** in the next tab.")
    st.session_state["report_md"] = st.text_area(
        label="Business Plan (Markdown)",
        value=st.session_state["report_md"],
        height=600,
    )
    st.download_button(
        "⬇️ Download Markdown",
        data=st.session_state["report_md"].encode("utf-8"),
        file_name="bizplan.md",
        mime="text/markdown",
        use_container_width=True
    )

# ========== TAB 4: Integrations & Export ==========
with tab4:
    st.subheader("Local API Integration (optional)")
    st.caption("If you have the FastAPI service from VS Code (your /generate endpoint), you can call it here.")
    colx1, colx2 = st.columns([3,1])
    with colx1:
        base_url = st.text_input("Base URL of your service", value="http://127.0.0.1:8000")
    with colx2:
        st.write(" ")
        if st.button("Call /generate", use_container_width=True):
            payload = {
                "prompt": user_prompt or default_prompt(),
                "assumptions": assump,
                "detail": "Investor Deep-Dive (60–100 pages)",
                "pages": 90
            }
            try:
                resp = call_local_api(base_url, payload)
                # Merge results (if API returns markdown/metrics)
                md = resp.get("markdown")
                if md:
                    st.session_state["report_md"] = md
                api_metrics = resp.get("metrics")
                if api_metrics:
                    st.info("Merged metrics from API into view (not overwriting model inputs).")
                    st.json(api_metrics)
                st.success("Local API call succeeded.")
            except Exception as e:
                st.error(f"Local API call failed: {e}")

    st.divider()
    st.subheader("Export")

    colz1, colz2, colz3, colz4 = st.columns(4)
    with colz1:
        st.download_button(
            "⬇️ Download Financials (CSV)",
            data=to_csv_bytes(df),
            file_name="financials.csv",
            mime="text/csv",
            use_container_width=True
        )
    with colz2:
        st.download_button(
            "⬇️ Download Debt Schedule (CSV)",
            data=to_csv_bytes(debt_df),
            file_name="debt_schedule.csv",
            mime="text/csv",
            use_container_width=True
        )
    with colz3:
        try:
            pdf_bytes = markdown_to_pdf_bytes(st.session_state["report_md"])
            st.download_button(
                "⬇️ Download PDF",
                data=pdf_bytes,
                file_name="bizplan.pdf",
                mime="application/pdf",
                use_container_width=True
            )
        except Exception as e:
            st.error(f"PDF generation error: {e}\nTry simplifying tables or keep to basic Markdown.")
    with colz4:
        try:
            docx_bytes = markdown_to_docx_bytes(st.session_state["report_md"])
            st.download_button(
                "⬇️ Download DOCX",
                data=docx_bytes,
                file_name="bizplan.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        except Exception as e:
            st.error(f"DOCX generation error: {e}")

    st.divider()
    st.subheader("Session")
    csa, csb = st.columns(2)
    with csa:
        if st.button("💾 Download Session (JSON)", use_container_width=True):
            blob = {
                "assumptions": assump,
                "metrics": metrics,
                "report_md": st.session_state["report_md"]
            }
            st.download_button(
                "Download session.json",
                data=json.dumps(blob, indent=2).encode("utf-8"),
                file_name="session.json",
                mime="application/json",
                use_container_width=True
            )
    with csb:
        uploaded = st.file_uploader("Restore Session (JSON)", type=["json"])
        if uploaded:
            try:
                data = json.loads(uploaded.read().decode("utf-8"))
                if "assumptions" in data:
                    st.info("Assumptions found in file. Please re-enter them in the sidebar to recalc.")
                if "report_md" in data:
                    st.session_state["report_md"] = data["report_md"]
                    st.success("Report restored from file.")
            except Exception as e:
                st.error(f"Failed to restore: {e}")

st.caption("Built with Streamlit. PDF engine uses xhtml2pdf; DOCX is a simple converter for headings/paragraphs.")
